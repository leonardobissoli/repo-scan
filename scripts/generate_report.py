#!/usr/bin/env python3
"""
repo-scan - Report generator.

Reads the JSON produced by scan_repo.py and generates:
  1) A detailed DOCX report.
  2) An interactive HTML dashboard (0-100 gauge, severity cards, filterable
     findings table, verdict).

Usage:
  python3 generate_report.py <scan.json> [--docx out.docx] [--html out.html]
"""

import html as _html
import json
import os
import sys

SEV_COLOR = {
    "CRITICAL": "C0152F",
    "HIGH": "E8590C",
    "MEDIUM": "F2A900",
    "LOW": "3B82F6",
    "INFO": "6B7280",
}
SEV_ORDER = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]

BAND_COLOR = {
    "SAFE": "16A34A",
    "LOW RISK": "65A30D",
    "MODERATE RISK": "F2A900",
    "HIGH RISK": "E8590C",
    "CRITICAL": "C0152F",
    "SELF-SCAN": "6B7280",
}

DEFAULT_OUT_DIR = os.path.join(os.getcwd(), "repo-scan-output")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def build_docx(data, out_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    meta = data["meta"]
    score = data["score"]
    verdict = data["verdict"]
    band = verdict["band"]
    is_self_scan = bool(data.get("self_scan"))
    band_rgb = RGBColor.from_string(BAND_COLOR.get(band, "6B7280"))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Repository Security Report")
    run.bold = True
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("repo-scan").italic = True

    doc.add_paragraph()

    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_self_scan:
        r = vp.add_run(f"raw score {score}/100")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = band_rgb
    else:
        r = vp.add_run(f"SCORE {score}/100")
        r.bold = True
        r.font.size = Pt(30)
        r.font.color.rgb = band_rgb

    vp2 = doc.add_paragraph()
    vp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = vp2.add_run(f"{band}  -  {verdict['action']}")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = band_rgb

    vp3 = doc.add_paragraph()
    vp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vp3.add_run(verdict["text"]).font.size = Pt(10.5)

    if is_self_scan:
        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        b1 = banner.add_run(
            "This is a self-scan of repo-scan's own source. "
            "The score and verdict are NOT a security signal about this "
            "repository. The numeric score reflects the scanner matching its "
            "own regex literals and the intentional test fixtures under "
            "test-samples/. For a step-by-step explanation and triage "
            "playbook, see references/interpreting-findings.md."
        )
        b1.font.size = Pt(9.5)
        b1.font.color.rgb = RGBColor.from_string("6B7280")
        b1.italic = True

    doc.add_paragraph()

    doc.add_heading("Scan summary", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    rows = [
        ("Source", str(meta.get("source"))),
        ("Commit", str(meta.get("commit"))),
        ("Date (UTC)", str(meta.get("scanned_at"))),
        ("Rules applied", str(meta.get("rules_count"))),
        ("Total files", str(data["stats"]["total_files"])),
        ("Executable files", str(data["stats"]["exec_file_count"])),
        ("Auto-exec points", str(len(data["auto_exec_points"]))),
    ]
    for k, v in rows:
        cells = tbl.add_row().cells
        cells[0].text = k
        cells[1].text = v
        cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("Findings by severity", level=1)
    sc = data["severity_counts"]
    stab = doc.add_table(rows=1, cols=2)
    stab.style = "Light Grid Accent 1"
    stab.rows[0].cells[0].text = "Severity"
    stab.rows[0].cells[1].text = "Count"
    for c in stab.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True
    for sev in SEV_ORDER:
        cells = stab.add_row().cells
        run = cells[0].paragraphs[0].add_run(sev)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(SEV_COLOR[sev])
        cells[1].text = str(sc.get(sev, 0))

    doc.add_heading("Risk by category", level=1)
    cat = data["category_summary"]
    if cat:
        ctab = doc.add_table(rows=1, cols=3)
        ctab.style = "Light Grid Accent 1"
        for i, h in enumerate(["Category", "Findings", "Max severity"]):
            ctab.rows[0].cells[i].text = h
            ctab.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for cname, info in sorted(cat.items(), key=lambda kv: SEV_ORDER.index(kv[1]["max_sev"])):
            cells = ctab.add_row().cells
            cells[0].text = cname
            cells[1].text = str(info["count"])
            mr = cells[2].paragraphs[0].add_run(info["max_sev"])
            mr.bold = True
            mr.font.color.rgb = RGBColor.from_string(SEV_COLOR[info["max_sev"]])
    else:
        doc.add_paragraph("No findings recorded.")

    doc.add_heading("Automatic execution points", level=1)
    doc.add_paragraph(
        "Files that may run code automatically (hooks, manifests, install "
        "scripts). These are the highest priority for manual review:"
    )
    if data["auto_exec_points"]:
        for p in data["auto_exec_points"]:
            doc.add_paragraph(p, style="List Bullet")
    else:
        doc.add_paragraph("No automatic execution points detected.")

    doc.add_heading("Detailed findings", level=1)
    findings = data["findings"]
    if not findings:
        doc.add_paragraph("No findings. Clean scan.")
    else:
        for idx, f in enumerate(findings, 1):
            h = doc.add_paragraph()
            hr = h.add_run(f"{idx}. [{f['severity']}] {f['rule_id']} - {f['category']}")
            hr.bold = True
            hr.font.size = Pt(11)
            hr.font.color.rgb = RGBColor.from_string(SEV_COLOR[f["severity"]])
            if f.get("likely_false_positive"):
                tag = h.add_run("  [likely false positive]")
                tag.bold = True
                tag.font.size = Pt(9)
                tag.font.color.rgb = RGBColor.from_string("6B7280")

            loc = doc.add_paragraph()
            loc.add_run("Location: ").bold = True
            loc.add_run(f"{f['file']}:{f['line']}")

            d = doc.add_paragraph()
            d.add_run("What it is: ").bold = True
            d.add_run(f["description"])

            sn = doc.add_paragraph()
            sn.add_run("Snippet: ").bold = True
            code = sn.add_run(f["snippet"])
            code.font.name = "Consolas"
            code.font.size = Pt(9)

            rec = doc.add_paragraph()
            rec.add_run("Recommendation: ").bold = True
            rec.add_run(f["recommendation"])
            doc.add_paragraph()

    doc.add_heading("Methodology and limitations", level=1)
    doc.add_paragraph(
        "This report is the result of STATIC analysis: the scanner reads and "
        "classifies text, it never executes code from the repository. The score "
        "starts at 100 and subtracts points by severity with diminishing returns "
        "(penalty = base x sqrt(count); CRITICAL 45, HIGH 18, MEDIUM 6, LOW 1.5)."
    )
    doc.add_paragraph(
        "Limitations: static analysis does not understand intent or dynamic "
        "flow. It can produce false positives (e.g. vulnerabilities planted in "
        "tests, which are auto-downgraded) and false negatives (heavily obfuscated "
        "code or malicious logic spread across files). A high score reduces, but "
        "does not eliminate, the need to manually review the auto-execution points. "
        "Always prefer official sources and pin the version/commit."
    )

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# HTML DASHBOARD
# ---------------------------------------------------------------------------


def build_html(data, out_path):
    score = data["score"]
    verdict = data["verdict"]
    band = verdict["band"]
    band_color = "#" + BAND_COLOR.get(band, "6B7280")
    meta = data["meta"]
    sc = data["severity_counts"]
    cats = data["category_summary"]
    findings = data["findings"]
    is_self_scan = bool(data.get("self_scan"))
    body_class = "self-scan" if is_self_scan else ""
    gauge_deg = int(180 * (score / 100))
    score_display = (
        f'raw score {score}<span style="font-size:13px;color:var(--muted)">/100</span>'
        if is_self_scan
        else f'{score}<span style="font-size:18px;color:var(--muted)">/100</span>'
    )
    self_scan_banner = (
        '<div class="self-scan-banner">'
        '<div class="ssb-label">SELF-SCAN DETECTED</div>'
        '<div class="ssb-text">This is a scan of repo-scan\'s own source. '
        "The score and verdict below are NOT a security signal about this "
        "repository - the scanner matches its own regex literals and the "
        "intentional <code>test-samples/</code> fixtures, which guarantees "
        "a CRITICAL raw score for any working copy of repo-scan. "
        '<a href="https://github.com/leonardobissoli/repo-scan/blob/main/references/interpreting-findings.md" '
        'target="_blank">How to read the findings &rarr;</a>'
        "</div></div>"
        if is_self_scan
        else ""
    )

    def esc(s):
        return _html.escape(str(s))

    sev_cards = "".join(
        f'<div class="sevcard" style="--c:#{SEV_COLOR[s]}">'
        f'<div class="sevn">{sc.get(s,0)}</div><div class="sevl">{s}</div></div>'
        for s in SEV_ORDER
    )

    cat_rows = (
        "".join(
            f'<tr><td>{esc(c)}</td><td style="text-align:center">{info["count"]}</td>'
            f'<td style="text-align:center"><span class="pill" style="background:#{SEV_COLOR[info["max_sev"]]}">{info["max_sev"]}</span></td></tr>'
            for c, info in sorted(cats.items(), key=lambda kv: SEV_ORDER.index(kv[1]["max_sev"]))
        )
        or '<tr><td colspan="3" style="text-align:center;color:#16A34A">No findings</td></tr>'
    )

    exec_items = (
        "".join(f"<li><code>{esc(p)}</code></li>" for p in data["auto_exec_points"])
        or "<li style='color:#16A34A'>No auto-execution points</li>"
    )

    def _fp_badge(f):
        return (
            '<span class="fp-badge" title="Match looks like documentation or '
            "scanner source rather than a live payload. Informational only - "
            'verify manually.">likely FP</span>'
            if f.get("likely_false_positive")
            else ""
        )

    find_rows = (
        "".join(
            f'<tr data-sev="{f["severity"]}" data-fp="{"1" if f.get("likely_false_positive") else "0"}">'
            f'<td><span class="pill" style="background:#{SEV_COLOR[f["severity"]]}">{f["severity"]}</span>{_fp_badge(f)}</td>'
            f'<td><b>{esc(f["rule_id"])}</b><br><span class="muted">{esc(f["category"])}</span></td>'
            f'<td><code>{esc(f["file"])}:{f["line"]}</code></td>'
            f'<td>{esc(f["description"])}<br><code class="snip">{esc(f["snippet"])}</code>'
            f'<div class="rec">&rarr; {esc(f["recommendation"])}</div></td>'
            f'</tr>'
            for f in findings
        )
        or '<tr><td colspan="4" style="text-align:center;color:#16A34A;padding:24px">Clean scan - no findings.</td></tr>'
    )

    fp_count = sum(1 for f in findings if f.get("likely_false_positive"))

    html = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>repo-scan - {esc(meta.get('source'))}</title>
<style>
:root{{--bg:#0b0e14;--card:#151a23;--line:#232b38;--txt:#e5e9f0;--muted:#8b94a7;}}
*{{box-sizing:border-box}}
body{{margin:0;background:var(--bg);color:var(--txt);font-family:Inter,-apple-system,Segoe UI,Roboto,sans-serif;padding:28px}}
h1{{font-size:20px;margin:0 0 2px}} .sub{{color:var(--muted);font-size:13px;margin-bottom:22px}}
.grid{{display:grid;grid-template-columns:320px 1fr;gap:18px;margin-bottom:18px}}
.card{{background:var(--card);border:1px solid var(--line);border-radius:14px;padding:20px}}
.gauge-wrap{{display:flex;flex-direction:column;align-items:center;justify-content:center}}
.gauge{{position:relative;width:240px;height:130px;overflow:hidden}}
.gauge .arc{{position:absolute;width:240px;height:240px;border-radius:50%;
 background:conic-gradient(from 270deg,#16A34A 0deg,#F2A900 90deg,#C0152F 180deg,transparent 180deg);
 -webkit-mask:radial-gradient(circle 76px at 120px 120px,transparent 75px,#000 76px);
 mask:radial-gradient(circle 76px at 120px 120px,transparent 75px,#000 76px);}}
.needle{{position:absolute;left:120px;bottom:0;width:3px;height:108px;background:var(--txt);
 transform-origin:bottom center;transform:rotate({gauge_deg-90}deg);border-radius:3px;z-index:3}}
.hub{{position:absolute;left:111px;bottom:-9px;width:18px;height:18px;border-radius:50%;background:var(--txt);z-index:4}}
.score{{font-size:46px;font-weight:800;line-height:1;margin-top:6px;color:{band_color}}}
.band{{font-weight:700;font-size:15px;color:{band_color};margin-top:4px}}
.action{{font-size:13px;color:var(--muted);text-align:center;margin-top:8px;max-width:260px}}
.sevrow{{display:flex;gap:10px;flex-wrap:wrap}}
.sevcard{{flex:1;min-width:80px;border:1px solid var(--line);border-left:4px solid var(--c);
 border-radius:10px;padding:12px 14px;background:#0f141d}}
.sevn{{font-size:26px;font-weight:800;color:var(--c)}} .sevl{{font-size:11px;color:var(--muted);letter-spacing:.5px}}
table{{width:100%;border-collapse:collapse;font-size:13px}}
th,td{{text-align:left;padding:9px 10px;border-bottom:1px solid var(--line);vertical-align:top}}
th{{color:var(--muted);font-weight:600;font-size:11px;text-transform:uppercase;letter-spacing:.5px}}
.pill{{color:#fff;padding:2px 9px;border-radius:20px;font-size:11px;font-weight:700}}
code{{font-family:ui-monospace,Menlo,Consolas,monospace;font-size:12px;color:#9ad}}
.snip{{display:block;margin-top:6px;color:#c5cdda;background:#0f141d;padding:6px 8px;border-radius:6px;white-space:pre-wrap;word-break:break-all}}
.rec{{margin-top:6px;color:#f2c97d;font-size:12px}}
.muted{{color:var(--muted);font-size:11px}}
.meta{{font-size:12px;color:var(--muted);line-height:1.7}}
.filters{{margin:6px 0 12px}} .filters button{{background:#0f141d;color:var(--txt);border:1px solid var(--line);
 padding:6px 12px;border-radius:20px;font-size:12px;cursor:pointer;margin-right:6px}}
.filters button.active{{background:{band_color};border-color:{band_color};color:#fff}}
.filters .sep{{display:inline-block;width:1px;height:20px;background:var(--line);vertical-align:middle;margin:0 6px}}
.fp-badge{{display:inline-block;margin-left:6px;padding:2px 7px;border-radius:10px;font-size:10px;font-weight:700;
 background:#1f2937;color:var(--muted);border:1px solid var(--line);cursor:help}}
tr[data-fp="1"]{{opacity:.55}} tr[data-fp="1"]:hover{{opacity:1}}
.hide-fp tr[data-fp="1"]{{display:none}}
.self-scan-banner{{background:#1f2937;border-left:5px solid #F2A900;border-radius:10px;padding:14px 18px;
 margin-bottom:18px;color:var(--txt)}}
.self-scan-banner .ssb-label{{font-size:11px;font-weight:800;letter-spacing:1.5px;color:#F2A900;margin-bottom:6px}}
.self-scan-banner .ssb-text{{font-size:13px;line-height:1.55;color:var(--txt)}}
.self-scan-banner a{{color:#9ad;text-decoration:underline;text-decoration-color:#374151}}
body.self-scan .gauge .arc{{background:#374151 !important}}
body.self-scan .score{{font-size:22px !important;color:#6B7280 !important}}
body.self-scan .band{{color:#6B7280 !important}}
.sec{{font-size:12px;text-transform:uppercase;letter-spacing:1px;color:var(--muted);margin:0 0 10px}}
.disclaimer{{font-size:11px;color:var(--muted);margin-top:18px;line-height:1.6}}
</style></head><body class="{body_class}">

<h1>Repository Security Report &middot; repo-scan</h1>
<div class="sub">{esc(meta.get('source'))} &nbsp;&middot;&nbsp; commit {esc(meta.get('commit'))} &nbsp;&middot;&nbsp; {esc(meta.get('scanned_at'))}</div>

{self_scan_banner}
<div class="grid">
  <div class="card gauge-wrap">
    <div class="gauge"><div class="arc"></div><div class="needle"></div><div class="hub"></div></div>
    <div class="score">{score_display}</div>
    <div class="band">{esc(band)}</div>
    <div class="action"><b>{esc(verdict['action'])}</b><br>{esc(verdict['text'])}</div>
  </div>
  <div class="card">
    <div class="sec">Findings by severity</div>
    <div class="sevrow">{sev_cards}</div>
    <div class="sec" style="margin-top:20px">Risk by category</div>
    <table><thead><tr><th>Category</th><th style="text-align:center">Findings</th><th style="text-align:center">Max sev.</th></tr></thead>
    <tbody>{cat_rows}</tbody></table>
  </div>
</div>

<div class="grid" style="grid-template-columns:320px 1fr">
  <div class="card">
    <div class="sec">Auto-execution points</div>
    <ul style="padding-left:18px;font-size:13px;line-height:1.8">{exec_items}</ul>
    <div class="sec" style="margin-top:18px">Statistics</div>
    <div class="meta">
      Total files: <b>{data['stats']['total_files']}</b><br>
      Executable: <b>{data['stats']['exec_file_count']}</b><br>
      Text analyzed: <b>{data['stats']['scanned_text_files']}</b><br>
      Rules applied: <b>{meta.get('rules_count')}</b>
    </div>
  </div>
  <div class="card">
    <div class="sec">Detailed findings</div>
    <div class="filters">
      <button class="active" onclick="flt(this,'ALL')">All</button>
      <button onclick="flt(this,'CRITICAL')">Critical</button>
      <button onclick="flt(this,'HIGH')">High</button>
      <button onclick="flt(this,'MEDIUM')">Medium</button>
      <button onclick="flt(this,'LOW')">Low</button>
      <span class="sep"></span>
      <button onclick="togFP(this)" title="{fp_count} finding(s) flagged as likely false positive">Hide likely FP ({fp_count})</button>
    </div>
    <table id="ftbl"><thead><tr><th>Sev</th><th>Rule</th><th>Location</th><th>Detail & Recommendation</th></tr></thead>
    <tbody id="ftab">{find_rows}</tbody></table>
  </div>
</div>

<div class="disclaimer">
<b>Static</b> analysis: the scanner reads and classifies text, it never executes the target's code.
Score starts at 100 and subtracts by severity (penalty = base x sqrt(count); CRITICAL 45, HIGH 18, MEDIUM 6, LOW 1.5).
False positives are possible (vulnerabilities planted in tests are downgraded) and false negatives too
(heavy obfuscation). A high score does not remove the need to manually review the auto-execution points.
Prefer official sources and pin the version/commit.
</div>

<script>
function flt(btn,sev){{
  document.querySelectorAll('.filters button').forEach((b,i)=>{{ if(i<5) b.classList.remove('active'); }});
  btn.classList.add('active');
  document.querySelectorAll('#ftab tr').forEach(tr=>{{
    tr.dataset.sevHidden = (sev!=='ALL' && tr.dataset.sev!==sev) ? '1':'0';
    tr.style.display = tr.dataset.sevHidden==='1' ? 'none' : '';
  }});
}}
function togFP(btn){{
  var tbl = document.getElementById('ftbl');
  tbl.classList.toggle('hide-fp');
  btn.classList.toggle('active');
}}
</script>
</body></html>"""

    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write(html)
    return out_path


def main():
    args = sys.argv[1:]
    if not args:
        print(
            "Usage: generate_report.py <scan.json> " "[--docx out.docx] [--html out.html] [--force]"
        )
        sys.exit(1)
    json_path = args[0]
    with open(json_path, encoding="utf-8") as fh:
        data = json.load(fh)

    force = "--force" in args
    if data.get("self_scan") and not force:
        print(
            "Input JSON is from a self-scan (target is a copy of repo-scan).\n"
            "Skipping DOCX and HTML generation - the verdict is "
            '"NOT REPRESENTATIVE" and a written report would be misleading.\n'
            "\n"
            "The raw findings are already in:\n"
            f"  {json_path}\n"
            "\n"
            "If you really need the artifacts (debug, screenshot of the banner, "
            "demo), re-run with --force."
        )
        return

    os.makedirs(DEFAULT_OUT_DIR, exist_ok=True)
    docx_out = os.path.join(DEFAULT_OUT_DIR, "repo_scan_report.docx")
    html_out = os.path.join(DEFAULT_OUT_DIR, "repo_scan_dashboard.html")
    if "--docx" in args:
        docx_out = args[args.index("--docx") + 1]
    if "--html" in args:
        html_out = args[args.index("--html") + 1]

    os.makedirs(os.path.dirname(docx_out) or ".", exist_ok=True)
    build_html(data, html_out)
    print(f"HTML: {html_out}")
    build_docx(data, docx_out)
    print(f"DOCX: {docx_out}")


if __name__ == "__main__":
    main()
